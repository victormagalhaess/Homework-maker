from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def msword(header, theme, data, picsLocation):
    document = Document()
    document.add_heading(theme, 0)
    document.add_heading(header, level=2)
    for paragraph in data:
        print(paragraph)
        if (paragraph != ''):
            if (paragraph[-1] in ['.', '!', '?', ':', ';']):
                document.add_paragraph(paragraph).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for picture in os.listdir(os.path.abspath(picsLocation)):
        document.add_picture(os.path.join(os.path.abspath(picsLocation), picture), width=Inches(4))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(theme +'_homework.docx')